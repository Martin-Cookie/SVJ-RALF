"""Generate personalized PDF ballot files for voting."""
import os
from typing import Optional

from docxtpl import DocxTemplate


def generate_ballot_pdf(
    template_path: str,
    output_path: str,
    voting_name: str,
    owner_name: str,
    unit_numbers: str,
    items: list,
) -> str:
    """Generate a personalized ballot document from a .docx template.

    Uses docxtpl to fill in template variables and saves as .docx.
    The output is a .docx file (PDF conversion requires LibreOffice which is optional).

    Returns the path to the generated file.
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    if template_path and os.path.exists(template_path):
        doc = DocxTemplate(template_path)
        context = {
            "voting_name": voting_name,
            "owner_name": owner_name,
            "unit_numbers": unit_numbers,
            "items": items,
        }
        doc.render(context)
        doc.save(output_path)
    else:
        # Fallback: create a simple document without template
        from docx import Document

        doc = Document()
        doc.add_heading(f"Hlasovací lístek — {voting_name}", level=1)
        doc.add_paragraph(f"Vlastník: {owner_name}")
        doc.add_paragraph(f"Jednotky: {unit_numbers}")
        doc.add_heading("Body hlasování", level=2)
        for i, item_text in enumerate(items, 1):
            doc.add_paragraph(f"{i}. {item_text}")
            doc.add_paragraph("   [ ] PRO    [ ] PROTI    [ ] Zdržel se")
        doc.save(output_path)

    return output_path
