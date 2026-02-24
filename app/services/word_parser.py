"""Parse voting items and metadata from .docx template files."""
import re
from typing import Dict, List, Optional

from docx import Document


def parse_voting_items(file_path: str) -> List[str]:
    """Extract numbered voting items from a .docx file.

    Looks for paragraphs starting with a number followed by a period or parenthesis,
    and returns the text of each item.
    """
    doc = Document(file_path)
    items = []
    pattern = re.compile(r"^\s*(\d+)\s*[.)]\s*(.+)")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        match = pattern.match(text)
        if match:
            items.append(match.group(2).strip())

    return items


_CZECH_MONTHS = {
    "ledna": 1, "února": 2, "března": 3, "dubna": 4,
    "května": 5, "června": 6, "července": 7, "srpna": 8,
    "září": 9, "října": 10, "listopadu": 11, "prosince": 12,
}


def _parse_czech_date(text: str) -> Optional[str]:
    """Try to parse a Czech date from text, return ISO YYYY-MM-DD or None."""
    # DD.MM.YYYY
    m = re.search(r"(\d{1,2})\.\s*(\d{1,2})\.\s*(\d{4})", text)
    if m:
        d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        return f"{y}-{mo:02d}-{d:02d}"
    # DD. měsíce YYYY  (e.g. "19. ledna 2026")
    m = re.search(
        r"(\d{1,2})\.\s*(" + "|".join(_CZECH_MONTHS) + r")\s+(\d{4})",
        text, re.IGNORECASE,
    )
    if m:
        d = int(m.group(1))
        mo = _CZECH_MONTHS[m.group(2).lower()]
        y = int(m.group(3))
        return f"{y}-{mo:02d}-{d:02d}"
    return None


def extract_voting_metadata(docx_path: str) -> Dict[str, Optional[str]]:
    """Extract voting metadata (name, dates) from a Word document.

    Returns dict with keys: name, start_date, end_date.
    Dates are in ISO format YYYY-MM-DD.
    """
    doc = Document(docx_path)
    paragraphs = [
        (p.text.strip(), p.style.name if p.style else "")
        for p in doc.paragraphs
    ]

    result: Dict[str, Optional[str]] = {
        "name": None,
        "start_date": None,
        "end_date": None,
    }

    # --- Name (title) ---
    per_rollam_pattern = re.compile(
        r"(hlasování\s+per\s+rollam|rozhodování\s+per\s+rollam|per\s+rollam)",
        re.IGNORECASE,
    )
    title_idx = None

    # 1. Document properties
    try:
        if doc.core_properties.title and doc.core_properties.title.strip():
            result["name"] = doc.core_properties.title.strip()
    except Exception:
        pass

    if not result["name"]:
        # 2. First Heading 1 or Title style
        for i, (text, style) in enumerate(paragraphs):
            if text and style in ("Heading 1", "Title"):
                result["name"] = text
                title_idx = i
                break

    if not result["name"]:
        # 3. Paragraph matching "per rollam" / "rozhodování per rollam"
        for i, (text, _) in enumerate(paragraphs):
            if text and per_rollam_pattern.search(text):
                result["name"] = text
                title_idx = i
                break

    # If title found and next paragraph is a short date-like line, merge it
    if result["name"] and title_idx is not None:
        next_idx = title_idx + 1
        while next_idx < len(paragraphs) and not paragraphs[next_idx][0]:
            next_idx += 1
        if next_idx < len(paragraphs):
            next_text = paragraphs[next_idx][0]
            if next_text and len(next_text) <= 80 and _parse_czech_date(next_text):
                result["name"] = result["name"] + " " + next_text
                title_idx = next_idx

    if not result["name"]:
        # 4. First short non-empty paragraph (max 150 chars)
        for i, (text, _) in enumerate(paragraphs):
            if text and len(text) <= 150:
                result["name"] = text
                title_idx = i
                break

    # --- Dates ---
    full_text = "\n".join(text for text, _ in paragraphs)

    start_patterns = [
        r"(?:od|zahájení|začátek)[:\s]+(\d{1,2}\.\s*\d{1,2}\.\s*\d{4})",
        r"(?:od|zahájení|začátek)[:\s]+(\d{1,2}\.\s*(?:" + "|".join(_CZECH_MONTHS) + r")\s+\d{4})",
        r"(?:vyhlášen[áéo])\s+(\d{1,2}\.\s*(?:" + "|".join(_CZECH_MONTHS) + r")\s+\d{4})",
        r"(?:vyhlášen[áéo])\s+(\d{1,2}\.\s*\d{1,2}\.\s*\d{4})",
    ]
    for pat in start_patterns:
        m = re.search(pat, full_text, re.IGNORECASE)
        if m:
            result["start_date"] = _parse_czech_date(m.group(1))
            if result["start_date"]:
                break

    end_patterns = [
        r"(?:do|ukončení|konec|odevzdání)[:\s]+.*?(\d{1,2}\.\s*\d{1,2}\.\s*\d{4})",
        r"(?:do|ukončení|konec|odevzdání)[:\s]+.*?(\d{1,2}\.\s*(?:" + "|".join(_CZECH_MONTHS) + r")\s+\d{4})",
        r"(?:lhůta|termín).*?(?:do|je)\s+(\d{1,2}\.\s*(?:" + "|".join(_CZECH_MONTHS) + r")\s+\d{4})",
        r"(?:lhůta|termín).*?(?:do|je)\s+(\d{1,2}\.\s*\d{1,2}\.\s*\d{4})",
    ]
    for pat in end_patterns:
        m = re.search(pat, full_text, re.IGNORECASE)
        if m:
            result["end_date"] = _parse_czech_date(m.group(1))
            if result["end_date"]:
                break

    return result
